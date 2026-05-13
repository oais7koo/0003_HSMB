"""Markdown to DOCX converter with proper heading styles"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def clean_markdown_text(text):
    """마크다운 문법 제거 (링크, 볼드 등)"""
    # [텍스트](링크) -> 텍스트
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    # **텍스트** -> 텍스트
    text = text.replace('**', '')
    return text

def make_bookmark_name(text):
    """텍스트를 북마크 이름으로 변환"""
    # 한글, 영문, 숫자만 유지하고 공백은 _로 변환
    name = re.sub(r'[^\w가-힣]', '_', text)
    name = re.sub(r'_+', '_', name).strip('_')
    return f"bm_{name[:30]}"  # 최대 30자

def add_bookmark(paragraph, bookmark_name):
    """문단에 북마크 추가"""
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)

    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')

    paragraph._p.insert(0, bookmark_start)
    paragraph._p.append(bookmark_end)

def add_hyperlink(paragraph, text, bookmark_name, indent=0):
    """내부 북마크로 연결되는 하이퍼링크 추가"""
    # 들여쓰기
    if indent > 0:
        paragraph.paragraph_format.left_indent = Cm(indent * 0.5)

    # 하이퍼링크 요소 생성
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Run 요소 생성
    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # 하이퍼링크 스타일 (파란색, 밑줄)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # 폰트 설정
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), '맑은 고딕')
    rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    rPr.append(rFonts)

    run.append(rPr)

    # 텍스트
    t = OxmlElement('w:t')
    t.text = text
    run.append(t)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)

def set_cell_text_with_breaks(cell, text):
    """셀에 텍스트 설정 (<br> 태그를 줄바꿈으로 변환)"""
    # 기존 텍스트 제거
    cell.text = ''
    paragraph = cell.paragraphs[0]

    # <br> 또는 <br/> 태그로 분할
    parts = re.split(r'<br\s*/?>', text)

    for i, part in enumerate(parts):
        run = paragraph.add_run(part.strip())
        run.font.name = '맑은 고딕'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        # 마지막 부분이 아니면 줄바꿈 추가
        if i < len(parts) - 1:
            run.add_break()

def set_table_borders(table):
    """표 전체 테두리 설정 (TableGrid 스타일)"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

    # 기존 스타일 제거
    for child in list(tblPr):
        if child.tag.endswith('tblStyle'):
            tblPr.remove(child)

    # TableGrid 스타일 (모든 셀에 테두리)
    tblStyle = OxmlElement('w:tblStyle')
    tblStyle.set(qn('w:val'), 'TableGrid')
    tblPr.insert(0, tblStyle)

    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # 추가로 명시적 테두리 설정
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)

def add_page_break(doc):
    """페이지 나누기 추가"""
    p = doc.add_paragraph()
    run = p.add_run()
    run._element.append(OxmlElement('w:br', {qn('w:type'): 'page'}))

def parse_markdown(md_content):
    """마크다운 파싱"""
    elements = []
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 빈 줄
        if not line.strip():
            i += 1
            continue

        # mermaid 블록 건너뛰기
        if line.strip().startswith('```mermaid'):
            while i < len(lines) and not lines[i].strip() == '```':
                i += 1
            i += 1
            elements.append({'type': 'skip', 'text': '[다이어그램 생략]'})
            continue

        # 코드 블록 건너뛰기
        if line.strip().startswith('```'):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                i += 1
            i += 1
            continue

        # 구분선
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # 제목
        heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = clean_markdown_text(heading_match.group(2))
            bookmark_name = make_bookmark_name(text)
            elements.append({'type': 'heading', 'level': level, 'text': text, 'bookmark': bookmark_name})
            i += 1
            continue

        # 표
        if line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # 구분선 건너뛰기
                if re.match(r'^\|[\s\-:|]+\|$', row_line):
                    i += 1
                    continue
                cells = [clean_markdown_text(c.strip()) for c in row_line.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1
            if table_rows:
                elements.append({'type': 'table', 'rows': table_rows})
            continue

        # 목록 (탭을 4칸 공백으로 변환하여 들여쓰기 계산)
        if re.match(r'^[-*]\s', line) or re.match(r'^[\s\t]+[-*]\s', line):
            raw_text = re.sub(r'^[\s\t]*[-*]\s+', '', line)
            # 탭을 4칸으로 변환하여 들여쓰기 레벨 계산
            line_expanded = line.replace('\t', '    ')
            indent = len(line_expanded) - len(line_expanded.lstrip())

            # 목차 링크 감지 [텍스트](#앵커)
            toc_match = re.match(r'^\[([^\]]+)\]\(#([^)]+)\)(.*)$', raw_text.strip())
            if toc_match:
                link_text = toc_match.group(1).replace('**', '')
                anchor = toc_match.group(2)
                suffix = clean_markdown_text(toc_match.group(3))
                # 앵커를 북마크 이름으로 변환
                bookmark_name = make_bookmark_name(link_text)
                elements.append({'type': 'toc_link', 'text': link_text, 'suffix': suffix,
                                'anchor': anchor, 'bookmark': bookmark_name, 'indent': indent})
            else:
                text = clean_markdown_text(raw_text)
                elements.append({'type': 'bullet', 'text': text, 'indent': indent})
            i += 1
            continue

        # 일반 문단
        text = clean_markdown_text(line).strip()
        if text:
            elements.append({'type': 'paragraph', 'text': text})
        i += 1

    return elements

def create_docx(elements, output_path):
    """DOCX 문서 생성"""
    doc = Document()

    # 기본 스타일 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    # 줄간격 20% 감소 (1.15 -> 0.92)
    style.paragraph_format.line_spacing = 0.92

    # 제목 스타일 설정
    for i in range(1, 5):
        try:
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = '맑은 고딕'
            heading_style.font.bold = True
            heading_style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
            if i == 1:
                heading_style.font.size = Pt(16)
            elif i == 2:
                heading_style.font.size = Pt(14)
            elif i == 3:
                heading_style.font.size = Pt(12)
            else:
                heading_style.font.size = Pt(11)
        except:
            pass

    title_added = False
    chapter_count = 0  # 장(##) 카운트

    for el in elements:
        if el['type'] == 'heading':
            level = el['level']
            bookmark_name = el.get('bookmark', '')

            if level == 1 and not title_added:
                # 문서 제목 추가 (중앙 정렬, 20pt)
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(el['text'])
                run.bold = True
                run.font.size = Pt(20)
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
                if bookmark_name:
                    add_bookmark(p, bookmark_name)
                doc.add_paragraph()  # 제목 뒤 여백
                title_added = True

            elif level == 2:
                # 장(##): 한 줄 비움 후 Heading 1
                if chapter_count > 0:
                    doc.add_paragraph()  # 장 전 한 줄 비움
                chapter_count += 1
                p = doc.add_heading(el['text'], level=1)
                if bookmark_name:
                    add_bookmark(p, bookmark_name)

            elif level == 3:
                # 절(###): 한 줄 비움 후 Heading 2
                doc.add_paragraph()  # 섹션 전 한 줄 비움
                p = doc.add_heading(el['text'], level=2)
                if bookmark_name:
                    add_bookmark(p, bookmark_name)

            elif level == 4:
                # 항(####): Heading 3
                p = doc.add_heading(el['text'], level=3)
                if bookmark_name:
                    add_bookmark(p, bookmark_name)

            elif level == 1:
                # 다른 # 제목
                p = doc.add_heading(el['text'], level=1)

            else:
                # 5, 6레벨
                p = doc.add_paragraph(el['text'])
                if p.runs:
                    p.runs[0].bold = True

        elif el['type'] == 'paragraph':
            doc.add_paragraph(el['text'])

        elif el['type'] == 'skip':
            p = doc.add_paragraph(el['text'])
            if p.runs:
                p.runs[0].italic = True

        elif el['type'] == 'toc_link':
            # 목차 링크 (하이퍼링크 포함)
            indent = el.get('indent', 0)
            indent_level = indent // 4  # 들여쓰기 레벨

            p = doc.add_paragraph()
            p.style = 'List Bullet'
            if indent_level > 0:
                p.paragraph_format.left_indent = Cm(indent_level * 0.5)

            # 하이퍼링크 추가
            add_hyperlink(p, el['text'], el['bookmark'], 0)

            # 접미사가 있으면 추가
            if el.get('suffix'):
                run = p.add_run(el['suffix'])
                run.font.name = '맑은 고딕'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

        elif el['type'] == 'bullet':
            # 들여쓰기에 따라 리스트 레벨 결정
            indent = el.get('indent', 0)
            if indent >= 4:
                style_name = 'List Bullet 3'
            elif indent >= 2:
                style_name = 'List Bullet 2'
            else:
                style_name = 'List Bullet'

            try:
                p = doc.add_paragraph(el['text'], style=style_name)
            except:
                # 스타일이 없으면 기본 스타일 + 수동 들여쓰기
                p = doc.add_paragraph(el['text'], style='List Bullet')
                p.paragraph_format.left_indent = Cm(indent * 0.5)

        elif el['type'] == 'table':
            rows = el['rows']
            if rows:
                max_cols = max(len(row) for row in rows)
                table = doc.add_table(rows=len(rows), cols=max_cols)

                # 표 테두리 스타일 적용
                set_table_borders(table)

                # 컬럼별 최대 글자 수 계산
                col_lengths = [0] * max_cols
                for row_data in rows:
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < max_cols:
                            col_lengths[col_idx] = max(col_lengths[col_idx], len(cell_text))

                # 컬럼 폭 비율 계산 및 적용 (전체 폭 16cm 기준)
                total_length = sum(col_lengths) or 1
                total_width = Cm(16)
                for col_idx, length in enumerate(col_lengths):
                    ratio = max(length / total_length, 0.05)  # 최소 5%
                    table.columns[col_idx].width = int(total_width * ratio)

                for row_idx, row_data in enumerate(rows):
                    for col_idx, cell_text in enumerate(row_data):
                        if col_idx < max_cols:
                            cell = table.rows[row_idx].cells[col_idx]
                            # <br> 태그를 줄바꿈으로 변환
                            set_cell_text_with_breaks(cell, cell_text)
                            # 첫 행 헤더 스타일
                            if row_idx == 0:
                                set_cell_shading(cell, 'D9E2F3')
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.bold = True

                doc.add_paragraph()  # 표 뒤 여백

    doc.save(output_path)
    print(f"Created: {output_path}")

def main():
    import sys

    if len(sys.argv) < 2 or sys.argv[1] in ['-h', '--help']:
        print("Usage: uv run .claude/skills/ccdoc/scripts/oodoc_md_to_docx.py <input.md> [output.docx]")
        print("       ccdoc run md <input.md>")
        print("")
        print("Markdown to DOCX 변환 (ccdoc 기본 스타일 적용)")
        print("  - 폰트: 맑은 고딕, 10pt")
        print("  - 줄간격: 0.92")
        print("  - 표: 전체 테두리, 헤더 배경색")
        print("  - 목록: 들여쓰기별 스타일")
        sys.exit(0 if sys.argv[1:] and sys.argv[1] in ['-h', '--help'] else 1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        print(f"Error: File not found: {md_path}")
        sys.exit(1)

    # 출력 경로: 지정하거나 입력 파일과 같은 위치에 .docx로 생성
    if len(sys.argv) >= 3:
        output_path = Path(sys.argv[2])
    else:
        output_path = md_path.with_suffix('.docx')

    content = md_path.read_text(encoding='utf-8')
    elements = parse_markdown(content)
    print(f"Parsed: {len(elements)} elements")

    create_docx(elements, str(output_path))

if __name__ == "__main__":
    main()
