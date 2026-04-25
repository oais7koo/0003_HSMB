"""Markdown to Word converter"""
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import re
import sys
import subprocess
import argparse
import tempfile
import os
from pathlib import Path


def convert_mermaid_to_png(mermaid_code: str, output_path: str) -> bool:
    """Convert mermaid code to PNG using mmdc"""
    with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False, encoding='utf-8') as f:
        f.write(mermaid_code)
        mmd_file = f.name

    try:
        cmd = f'mmdc -i "{mmd_file}" -o "{output_path}" -b white'
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            shell=True
        )
        return result.returncode == 0
    finally:
        os.unlink(mmd_file)


def preprocess_mermaid(content: str, base_dir: str) -> tuple[str, list[str]]:
    """Extract mermaid blocks and convert to PNG images

    Returns:
        tuple: (processed_content, list_of_image_paths)
    """
    mermaid_pattern = r'```mermaid\n(.*?)```'
    matches = list(re.finditer(mermaid_pattern, content, re.DOTALL))

    if not matches:
        return content, []

    image_paths = []
    processed_content = content

    # Process in reverse order to maintain correct positions
    for i, match in enumerate(reversed(matches)):
        idx = len(matches) - 1 - i
        mermaid_code = match.group(1)

        # Create image path
        img_filename = f'_mermaid_{idx}.png'
        img_path = os.path.join(base_dir, img_filename)

        if convert_mermaid_to_png(mermaid_code, img_path):
            image_paths.append(img_path)
            # Replace mermaid block with image reference
            img_md = f'![Diagram {idx + 1}]({img_filename})'
            processed_content = (
                processed_content[:match.start()] +
                img_md +
                processed_content[match.end():]
            )
        else:
            print(f'Warning: mermaid 변환 실패 (block {idx + 1})')

    return processed_content, image_paths


def cleanup_mermaid_images(image_paths: list[str]):
    """Clean up temporary mermaid PNG files"""
    for path in image_paths:
        if os.path.exists(path):
            os.unlink(path)




def add_formatted_text(paragraph, text, default_font='Malgun Gothic', default_size=Pt(11)):
    pattern = r'(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
    last_end = 0
    for match in re.finditer(pattern, text):
        if match.start() > last_end:
            run = paragraph.add_run(text[last_end:match.start()])
            run.font.name = default_font
            run.font.size = default_size
        full_match = match.group(0)
        if full_match.startswith('**'):
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.font.name = default_font
            run.font.size = default_size
        elif full_match.startswith('`'):
            run = paragraph.add_run(match.group(4))
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif full_match.startswith('*'):
            run = paragraph.add_run(match.group(3))
            run.italic = True
            run.font.name = default_font
            run.font.size = default_size
        last_end = match.end()
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        run.font.name = default_font
        run.font.size = default_size

def set_cell_border(cell, border_color="000000", border_size="4"):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:left w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:bottom w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'<w:right w:val="single" w:sz="{border_size}" w:color="{border_color}"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)


def create_table(doc, table_rows):
    """Create Word table from markdown table rows"""
    if not table_rows:
        return

    # Parse cells from each row
    parsed_rows = []
    for row in table_rows:
        cells = [c.strip() for c in row.split('|')[1:-1]]
        # Skip separator row (|---|---|)
        if cells and not all(re.match(r'^:?-+:?$', c) for c in cells):
            parsed_rows.append(cells)

    if not parsed_rows:
        return

    # Determine table dimensions
    num_cols = max(len(row) for row in parsed_rows)
    num_rows = len(parsed_rows)

    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Fill cells
    for i, row_data in enumerate(parsed_rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]

                # Use add_formatted_text to handle **bold**, *italic*, `code`
                paragraph = cell.paragraphs[0]
                paragraph.clear()
                add_formatted_text(paragraph, cell_text, default_size=Pt(10))

                # Set border
                set_cell_border(cell)

                # Style: header row (first row) bold + center
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Set font for all runs
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = 'Malgun Gothic'
                        run.font.size = Pt(10)

    # No empty paragraph after table (removed)


def convert_md_to_docx(input_path: str, output_path: str):
    """Convert markdown file to Word document"""

    # Read markdown file
    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create Word document
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Malgun Gothic'
    style.font.size = Pt(11)

    # Parse and convert markdown to Word
    lines = content.split('\n')
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []

    for line in lines:
        # Code block handling
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_content)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.style = 'Normal'
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_content = []
            in_code_block = not in_code_block
            continue

        if in_code_block:
            code_content.append(line)
            continue

        # Table handling
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            continue
        elif in_table:
            # End of table
            create_table(doc, table_rows)
            in_table = False
            table_rows = []
            # Continue processing current line

        # Headers
        if line.startswith('# '):
            doc.add_heading(line[2:], level=0)
            doc.add_paragraph()  # 최상위 제목 뒤에만 빈 줄
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=3)
        elif line.startswith('##### '):
            doc.add_heading(line[6:], level=4)
        elif line.startswith('- '):
            # Bullet list
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, line[2:])
        elif re.match(r'^\d+\. ', line):
            # Numbered list
            text = re.sub(r'^\d+\. ', '', line)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
        elif line.startswith('> '):
            # Quote
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            add_formatted_text(p, line[2:])
        elif line.strip() == '':
            # Empty line - skip (no paragraph break)
            continue
        elif line.startswith('---'):
            # Horizontal rule - skip
            continue
        else:
            # Normal paragraph
            if line.strip():
                p = doc.add_paragraph()
                add_formatted_text(p, line)

    # Handle table at end of file
    if in_table and table_rows:
        create_table(doc, table_rows)

    # Save
    doc.save(output_path)
    print(f'Word 문서 생성 완료: {output_path}')


def convert_with_pandoc(input_path: str, output_path: str):
    """Convert markdown to Word using pandoc (supports LaTeX math and mermaid)"""
    # Use absolute paths
    input_abs = os.path.abspath(input_path)
    output_abs = os.path.abspath(output_path)
    base_dir = os.path.dirname(input_abs)

    # Read original content
    with open(input_abs, 'r', encoding='utf-8') as f:
        content = f.read()

    image_paths = []
    temp_md = None

    try:
        # Preprocess mermaid diagrams
        processed_content, image_paths = preprocess_mermaid(content, base_dir)

        # If mermaid was processed, use temp file
        if image_paths:
            temp_md = os.path.join(base_dir, '_temp_processed.md')
            with open(temp_md, 'w', encoding='utf-8') as f:
                f.write(processed_content)
            source_file = temp_md
            print(f'mermaid 다이어그램 {len(image_paths)}개 변환 완료')
        else:
            source_file = input_abs

        # Run pandoc with shell command string (Windows compatibility)
        cmd = f'pandoc "{source_file}" -o "{output_abs}" --resource-path="{base_dir}"'
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            shell=True
        )

        if result.returncode == 0:
            print(f'Word 문서 생성 완료 (pandoc): {output_path}')
        else:
            print(f'pandoc 오류: {result.stderr}')
            sys.exit(1)

    except FileNotFoundError:
        print('pandoc이 설치되어 있지 않습니다. pandoc을 설치해주세요.')
        sys.exit(1)
    finally:
        # Cleanup temp files
        if temp_md and os.path.exists(temp_md):
            os.unlink(temp_md)
        cleanup_mermaid_images(image_paths)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Markdown to Word converter')
    parser.add_argument('input_file', nargs='?', default='doc/d0102_가상논문.md',
                        help='Input markdown file')
    parser.add_argument('--pandoc', action='store_true',
                        help='Use pandoc for conversion (supports LaTeX math)')
    args = parser.parse_args()

    input_file = args.input_file
    output_file = input_file.rsplit('.', 1)[0] + '.docx'

    if args.pandoc:
        convert_with_pandoc(input_file, output_file)
    else:
        convert_md_to_docx(input_file, output_file)
